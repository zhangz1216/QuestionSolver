"""题库管理：多格式导入（PDF/Word/TXT/图片）-> 文本库 -> BM25 检索。

存储：%LOCALAPPDATA%/QuestionSolver/data.db（SQLite）
检索：jieba 分词 + 简化 BM25，纯本地计算，不花 API 钱。
"""
import os
import re
import sqlite3
import math
import time
from pathlib import Path

import jieba

from . import ocr

DATA_DIR = Path(os.environ.get("LOCALAPPDATA", Path.home())) / "QuestionSolver"
DATA_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = DATA_DIR / "data.db"

CHUNK_SIZE = 600      # 分块字符数
CHUNK_OVERLAP = 80    # 相邻块重叠，避免跨块断句
TOP_K = 3             # 检索返回最多块数
PDF_MIN_CHARS_PER_PAGE = 30  # 低于此字数视为扫描版 PDF，转图片 OCR

KINDS = {
    "txt": "文本",
    "pdf": "PDF",
    "docx": "Word",
    "image": "图片",
}

_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".webp"}


# ---------------------------------------------------------------- 解析
def _read_text_file(path: str) -> str:
    data = Path(path).read_bytes()
    for enc in ("utf-8", "gbk", "utf-16"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, ValueError):
            continue
    return data.decode("utf-8", errors="replace")


def extract_text_pdf(path: str) -> str:
    """PDF 提取文本。扫描版（无文字层）自动转图片 OCR。"""
    import pymupdf

    doc = pymupdf.open(path)
    parts = []
    for page in doc:
        text = page.get_text("text").strip()
        if len(text) < PDF_MIN_CHARS_PER_PAGE:
            # 扫描版：渲染成图片 OCR
            pix = page.get_pixmap(dpi=200)
            png_bytes = pix.tobytes("png")
            page_text = ocr.recognize(png_bytes).strip()
            if page_text:
                parts.append(page_text)
        else:
            parts.append(text)
    doc.close()
    return "\n".join(parts)


def extract_text_docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n".join(paras)


def extract_text(path: str) -> str:
    """按扩展名解析文件为纯文本。"""
    p = Path(path)
    ext = p.suffix.lower()
    if ext in {".txt", ".md"}:
        return _read_text_file(str(p))
    if ext == ".pdf":
        return extract_text_pdf(str(p))
    if ext == ".docx":
        return extract_text_docx(str(p))
    if ext in _IMAGE_EXTS:
        return ocr.recognize(p.read_bytes()).strip()
    raise ValueError(f"不支持的文件类型：{ext}（支持 TXT/MD/PDF/DOCX/图片）")


def kind_of(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext in {".txt", ".md"}:
        return "txt"
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext in _IMAGE_EXTS:
        return "image"
    return ""


# ---------------------------------------------------------------- 存储
def _connect():
    conn = sqlite3.connect(str(DB_PATH))
    conn.execute("""CREATE TABLE IF NOT EXISTS questionbank (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        kind TEXT NOT NULL,
        text TEXT NOT NULL,
        char_count INTEGER DEFAULT 0,
        imported_at REAL NOT NULL
    )""")
    return conn


def import_file(path: str) -> dict:
    """导入一个文件到题库，返回条目 dict。"""
    name = Path(path).name
    kind = kind_of(path)
    if not kind:
        raise ValueError(f"不支持的文件类型：{Path(path).suffix}（支持 TXT/MD/PDF/DOCX/图片）")
    text = extract_text(path)
    if not text or len(text.strip()) < 10:
        raise ValueError(f"未从「{name}」提取到有效文本（可能是纯图片或已损坏）")
    conn = _connect()
    try:
        cur = conn.execute(
            "INSERT INTO questionbank (name, kind, text, char_count, imported_at) VALUES (?,?,?,?,?)",
            (name, kind, text, len(text), time.time()))
        conn.commit()
        row = {"id": cur.lastrowid, "name": name, "kind": kind,
               "char_count": len(text), "imported_at": time.time()}
        return row
    finally:
        conn.close()


def list_bank() -> list:
    conn = _connect()
    try:
        rows = conn.execute(
            "SELECT id, name, kind, char_count, imported_at FROM questionbank "
            "ORDER BY imported_at DESC").fetchall()
        return [{"id": r[0], "name": r[1], "kind": r[2],
                 "char_count": r[3], "imported_at": r[4]} for r in rows]
    finally:
        conn.close()


def delete_bank(bank_id: int):
    conn = _connect()
    try:
        conn.execute("DELETE FROM questionbank WHERE id=?", (bank_id,))
        conn.commit()
    finally:
        conn.close()


def get_bank_text(bank_id: int) -> str:
    conn = _connect()
    try:
        row = conn.execute("SELECT text FROM questionbank WHERE id=?", (bank_id,)).fetchone()
        return row[0] if row else ""
    finally:
        conn.close()


# ---------------------------------------------------------------- 检索
def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list:
    """把长文本切成块（保留语义边界：优先在段落/标点处切）。"""
    text = re.sub(r"\n{2,}", "\n", text).strip()
    if not text:
        return []
    # 先按段落分，再按长度合并/切分
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks, buf = [], ""
    for para in paragraphs:
        if len(buf) + len(para) <= size:
            buf += ("\n" + para) if buf else para
            continue
        if buf:
            chunks.append(buf)
            buf = ""
        # 超长段落直接硬切
        while len(para) > size:
            chunks.append(para[:size])
            para = para[size - overlap:]
        buf = para
    if buf:
        chunks.append(buf)
    return chunks


class _BM25Index:
    """简化 BM25：题库量不大，每次检索实时构建。"""

    def __init__(self, docs: list):
        self.docs = docs
        self.k1, self.b = 1.5, 0.75
        self.doc_freqs, self.doc_lens, self.df, self.avgdl = self._build(docs)

    def _build(self, docs):
        doc_freqs, doc_lens, df = [], [], {}
        for d in docs:
            tokens = list(jieba.cut(d))
            doc_lens.append(len(tokens))
            freq = {}
            for t in tokens:
                freq[t] = freq.get(t, 0) + 1
            doc_freqs.append(freq)
            for t in set(freq):
                df[t] = df.get(t, 0) + 1
        avgdl = sum(doc_lens) / len(doc_lens) if doc_lens else 1.0
        return doc_freqs, doc_lens, df, avgdl

    def score(self, query_tokens, doc_idx):
        freq, dl = self.doc_freqs[doc_idx], self.doc_lens[doc_idx]
        s = 0.0
        for t in query_tokens:
            tf = freq.get(t, 0)
            if not tf:
                continue
            idf = math.log((len(self.docs) - self.df.get(t, 0) + 0.5) /
                           (self.df.get(t, 0) + 0.5) + 1.0)
            s += idf * (tf * (self.k1 + 1)) / (tf + self.k1 * (1 - self.b + self.b * dl / self.avgdl))
        return s


def retrieve(question: str, bank_ids=None, top_k: int = TOP_K) -> list:
    """检索题库中与题目最相关的文本块。

    bank_ids=None 表示检索全部题库；否则只检索指定 id 列表。
    返回 [{source_id, source_name, text, score}]，按分数降序。
    """
    conn = _connect()
    try:
        if bank_ids:
            marks = ",".join("?" for _ in bank_ids)
            rows = conn.execute(
                f"SELECT id, name, text FROM questionbank WHERE id IN ({marks})",
                bank_ids).fetchall()
        else:
            rows = conn.execute(
                "SELECT id, name, text FROM questionbank").fetchall()
    finally:
        conn.close()

    if not rows:
        return []

    # 每份资料分块，记录来源
    docs, sources = [], []
    for rid, name, text in rows:
        for chunk in chunk_text(text):
            docs.append(chunk)
            sources.append((rid, name))

    query_tokens = list(jieba.cut(question))
    index = _BM25Index(docs)
    scored = sorted(range(len(docs)), key=lambda i: index.score(query_tokens, i), reverse=True)

    results = []
    seen = set()
    for i in scored:
        score = index.score(query_tokens, i)
        if score <= 0:
            break
        rid, name = sources[i]
        if rid in seen and len(results) >= top_k:
            continue
        seen.add(rid)
        results.append({"source_id": rid, "source_name": name,
                        "text": docs[i], "score": round(score, 3)})
        if len(results) >= top_k:
            break
    return results
